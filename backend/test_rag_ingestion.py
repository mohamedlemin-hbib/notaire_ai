from docx import Document
import requests
import json
import os

def create_sample_docx(path):
    doc = Document()
    doc.add_heading('Modèle d\'Acte de Vente Immobilière', 0)
    doc.add_paragraph('Ceci est un modèle standard pour une vente à Toulouse.')
    doc.add_paragraph('Clause de prix : Le prix est fixé à 250 000 euros.')
    doc.add_paragraph('Clause de condition suspensive : Obtention d\'un prêt immobilier.')
    doc.save(path)
    print(f"Fichier échantillon créé : {path}")

def test_upload_flow():
    base_url = "http://127.0.0.1:8002" # Match active server
    sample_path = "sample_vente.docx"
    create_sample_docx(sample_path)

    # 1. Login to get token
    print("Connexion en cours...")
    login_data = {"username": "mohamedleminaidelha@gmail.com", "password": "Notaire2024!"}
    res_login = requests.post(f"{base_url}/api/v1/auth/login", data=login_data)
    token = res_login.json()["access_token"]
    print("Token obtenu.")

    # 2. Upload template
    print("Upload du document...")
    headers = {"Authorization": f"Bearer {token}"}
    files = {'file': open(sample_path, 'rb')}
    params = {'act_type': 'vente'}
    
    res_upload = requests.post(
        f"{base_url}/api/v1/admin/upload-template",
        headers=headers,
        files=files,
        params=params
    )
    
    print(f"Statut Upload: {res_upload.status_code}")
    print(f"Réponse: {json.dumps(res_upload.json(), indent=2)}")

    # Cleanup
    os.unlink(sample_path)

if __name__ == "__main__":
    # Note: Make sure the server is running on 8001
    # Run: uvicorn app.main:app --host 127.0.0.1 --port 8001
    try:
        test_upload_flow()
    except Exception as e:
        print(f"Erreur: {e}")
